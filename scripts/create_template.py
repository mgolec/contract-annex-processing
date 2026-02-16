"""Build the aneks_template.docx programmatically using python-docx.

This creates a docxtpl-compatible template that mirrors the Crowe annex structure
(U-25-09) with Jinja2 placeholders for variable content.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates" / "default"
TEMPLATE_PATH = TEMPLATE_DIR / "aneks_template.docx"


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 10) -> None:
    """Set cell text with consistent formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold


def _add_para(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    space_after: int | None = None,
    space_before: int | None = None,
    font_size: int = 10,
) -> None:
    """Add a paragraph with consistent formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.bold = bold
    if alignment is not None:
        p.alignment = alignment
    fmt = p.paragraph_format
    if space_after is not None:
        fmt.space_after = Pt(space_after)
    if space_before is not None:
        fmt.space_before = Pt(space_before)


def _add_mixed_para(
    doc: Document,
    segments: list[tuple[str, bool]],
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    space_after: int | None = None,
    space_before: int | None = None,
) -> None:
    """Add a paragraph with mixed bold/normal runs."""
    p = doc.add_paragraph()
    for text, bold in segments:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.bold = bold
    if alignment is not None:
        p.alignment = alignment
    fmt = p.paragraph_format
    if space_after is not None:
        fmt.space_after = Pt(space_after)
    if space_before is not None:
        fmt.space_before = Pt(space_before)


def build_template() -> Path:
    """Create the annex .docx template with Jinja2 placeholders."""
    doc = Document()

    # ── Page setup (A4, 2cm margins) ────────────────────────────────
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # ── Set default font ────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # ── Header block: Client party ──────────────────────────────────
    _add_mixed_para(
        doc,
        [
            ("{{ korisnik_naziv }}", True),
            (", {{ korisnik_adresa }}", True),
            (", OIB: {{ korisnik_oib }}, kojeg zastupa direktor {{ korisnik_direktor }} "
             "(u daljnjem tekstu: Korisnik usluga),", False),
        ],
    )

    _add_para(doc, "")  # blank line

    _add_para(doc, "i", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    _add_para(doc, "")  # blank line

    # ── Header block: Procudo party ─────────────────────────────────
    _add_mixed_para(
        doc,
        [
            ("{{ davatelj_naziv }}", True),
            (" {{ davatelj_adresa }}, OIB: {{ davatelj_oib }}, kojeg zastupa direktor "
             "{{ davatelj_direktor }}, (u daljnjem tekstu: Izvršitelj usluga),", False),
        ],
    )

    _add_para(doc, "")  # blank line

    _add_para(doc, "zaključili su dana {{ datum_aneksa }} godine.")

    _add_para(doc, "")  # blank line

    # ── Title block ─────────────────────────────────────────────────
    _add_para(
        doc,
        "Anex br. {{ broj_aneksa }} Ugovora br. {{ broj_ugovora }}",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    _add_para(
        doc,
        "o servisiranju i održavanju",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    _add_para(
        doc,
        "informacijskog sustava",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )

    # ── Čl. 1 ───────────────────────────────────────────────────────
    _add_para(doc, "Čl. 1", bold=True, space_before=8)
    _add_para(doc, "")
    _add_para(
        doc,
        "Ugovorne strane suglasno utvrđuju da su dana {{ datum_ugovora }} "
        "sklopile Ugovor br. {{ broj_ugovora }} o servisiranju i održavanju "
        "informacijskog sustava",
    )

    # ── Čl. 2 ───────────────────────────────────────────────────────
    _add_para(doc, "Čl. 2", bold=True, space_before=8)
    _add_para(doc, "")
    _add_para(
        doc,
        "Ugovorne strane suglasno utvrđuju da se mijenja odredba čl. 3. "
        "gore navedenog Ugovora na način da ista sada glasi:",
    )
    _add_para(doc, "")
    _add_para(doc, "    Naknada mjesečnog održavanja", bold=True)
    _add_para(doc, "")
    # Conditional HRK conversion clause
    _add_para(
        doc,
        "{% if valuta_konverzija %}"
        "Dosadašnje cijene usluga bile su izražene u HRK te se konvertiraju u EUR "
        "prema fiksnom tečaju konverzije (1 EUR = 7,53450 HRK) sukladno Zakonu o "
        "uvođenju eura kao službene valute u Republici Hrvatskoj."
        "{% endif %}"
    )
    _add_para(
        doc,
        "Za usluge koje su definirane u Prilogu 1. ovog Ugovora, Korisnik usluga "
        "će Izvršitelju usluga plaćati mjesečnu naknadu u iznosu od "
        "{{ mjesecna_naknada }} EUR neto.",
    )

    # ── Čl. 3 ───────────────────────────────────────────────────────
    _add_para(doc, "Čl. 3", bold=True, space_before=8)
    _add_para(doc, "")
    _add_para(
        doc,
        "Ugovorne strane suglasno utvrđuju da se mijenja odredba čl. 4 st. 4 "
        "gore navedenog Ugovora na način da ista sada glasi:",
    )
    _add_para(doc, "")
    _add_para(doc, "    Obveze Izvršitelja usluga", bold=True)
    _add_para(doc, "")
    _add_mixed_para(
        doc,
        [
            ("(4) Ugovoreni mjesečni fond sati za redovito servisiranje i "
             "održavanje informacijskog sustava je ", False),
            ("{{ ukupno_sati }}", True),
            (" mjesečno:", False),
        ],
    )
    _add_para(doc, "")
    _add_para(
        doc,
        "{{ l1_sati }} sistem administrator sata (L1) – sistemsko održavanje "
        "radnih stanica, mrežne infrastrukture, printera i hardverske periferije, "
        "pomoć korisnicima pri radu, edukacija korisnika",
    )
    _add_para(doc, "")
    _add_para(
        doc,
        "{{ l2_sati }} sistem inženjer sat – (L2) – Održavanje i konfiguriranje "
        "komunikacijske mreže, održavanje i konfiguriranje poslužitelja i Office 365 "
        "sustava, savjetodavne usluge",
    )

    # ── Čl. 4 ───────────────────────────────────────────────────────
    _add_para(doc, "Čl. 4.", bold=True, space_before=8)
    _add_para(doc, "")
    _add_para(
        doc,
        "Ugovorne strane suglasno utvrđuju da se mijenja odredba Priloga 2. "
        "gore navedenog Ugovora na način da ista sada glasi:",
    )
    _add_para(doc, "")
    _add_para(
        doc,
        "Izračun cijene redovnog servisiranja i održavanja informacijskog sustava",
        bold=True,
    )
    _add_para(doc, "")

    # ── Pricing table ───────────────────────────────────────────────
    # docxtpl {%tr %} tags must be on their own dedicated rows.
    # Structure: header | for-loop-tag row | data row | endfor-tag row
    table = doc.add_table(rows=4, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Row 0: Header row
    headers = ["Poz.", "Opis", "Oznaka", "Mjera", "Kol.", "Jed. Cijena\n(EUR)"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, h, bold=True)

    # Row 1: for-loop opening tag (dedicated row)
    _set_cell_text(table.rows[1].cells[0], "{%tr for s in stavke %}")
    for i in range(1, 6):
        _set_cell_text(table.rows[1].cells[i], "")

    # Row 2: data row with placeholders
    data_cells = [
        "{{ s.pozicija }}",
        "{{ s.opis }}",
        "{{ s.oznaka }}",
        "{{ s.mjera }}",
        "{{ s.kolicina }}",
        "{{ s.cijena }}",
    ]
    for i, val in enumerate(data_cells):
        _set_cell_text(table.rows[2].cells[i], val)

    # Row 3: endfor tag (dedicated row)
    _set_cell_text(table.rows[3].cells[0], "{%tr endfor %}")
    for i in range(1, 6):
        _set_cell_text(table.rows[3].cells[i], "")

    _add_para(doc, "")
    _add_para(doc, "{{ vat_note }}")

    # ── Čl. 5 ───────────────────────────────────────────────────────
    _add_para(doc, "Čl. 5.", bold=True, space_before=8)
    _add_para(doc, "")
    _add_para(
        doc,
        "Ugovorne strane suglasno utvrđuju da ostale odredbe Ugovora "
        "{{ broj_ugovora }} o servisiranju i održavanju informacijskog "
        "sustava ostaju nepromijenjene.",
    )

    # ── Čl. 6 ───────────────────────────────────────────────────────
    _add_para(doc, "Čl. 6.", bold=True, space_before=8)
    _add_para(doc, "")
    _add_para(
        doc,
        "Ovaj Aneks je sastavljen u 2 istovjetna primjerka, od kojih svaka "
        "ugovorna strana zadržava po 1 primjerak.",
    )

    # ── Čl. 7 ───────────────────────────────────────────────────────
    _add_para(doc, "Čl. 7.", bold=True, space_before=8)
    _add_para(doc, "")
    _add_para(
        doc,
        "Korisnik usluga i Izvršitelj usluga prihvaćaju prava i obveze "
        "iz ovog Aneksa, te ga u znak obostranog prihvata vlastoručno "
        "potpisuju po ovlaštenim zastupnicima.",
    )

    _add_para(doc, "")

    # ── Signature block ─────────────────────────────────────────────
    _add_para(doc, "U {{ mjesto }}, {{ datum_aneksa }}")
    _add_para(doc, "")
    _add_para(doc, "")

    # Create a 2-column table for signature alignment (no borders)
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove borders from signature table
    for row in sig_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn("w:tcBorders"))
            if tcBorders is not None:
                tcPr.remove(tcBorders)
            # Add empty borders element to suppress
            borders = tcPr.makeelement(qn("w:tcBorders"), {})
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = borders.makeelement(qn(f"w:{edge}"), {
                    qn("w:val"): "none",
                    qn("w:sz"): "0",
                    qn("w:space"): "0",
                    qn("w:color"): "auto",
                })
                borders.append(border)
            tcPr.append(borders)

    # Row 0: Labels
    _set_cell_text(sig_table.rows[0].cells[0], "Korisnik usluga")
    _set_cell_text(sig_table.rows[0].cells[1], "Izvršitelj usluga")

    # Row 1: Company names
    _set_cell_text(sig_table.rows[1].cells[0], "{{ korisnik_naziv }}", bold=True)
    _set_cell_text(sig_table.rows[1].cells[1], "{{ davatelj_naziv }}", bold=True)

    # Row 2: Signature lines
    _set_cell_text(sig_table.rows[2].cells[0], "_________________")
    _set_cell_text(sig_table.rows[2].cells[1], "_________________")

    # Row 3: Director names + role
    _set_cell_text(sig_table.rows[3].cells[0], "{{ korisnik_direktor }}\nDirektor")
    _set_cell_text(sig_table.rows[3].cells[1], "{{ davatelj_direktor }}\nDirektor")

    # MP. (stamp) line
    _add_para(doc, "")
    sig2 = doc.add_table(rows=1, cols=2)
    sig2.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove borders
    for cell in sig2.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = tcPr.makeelement(qn("w:tcBorders"), {})
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = borders.makeelement(qn(f"w:{edge}"), {
                qn("w:val"): "none",
                qn("w:sz"): "0",
                qn("w:space"): "0",
                qn("w:color"): "auto",
            })
            borders.append(border)
        tcPr.append(borders)
    _set_cell_text(sig2.rows[0].cells[0], "MP.")
    _set_cell_text(sig2.rows[0].cells[1], "MP.")

    # ── Save ────────────────────────────────────────────────────────
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(TEMPLATE_PATH))
    print(f"Template created: {TEMPLATE_PATH}")
    return TEMPLATE_PATH


if __name__ == "__main__":
    build_template()
